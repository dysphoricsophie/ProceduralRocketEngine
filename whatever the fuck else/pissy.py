from time import sleep
from docx import Document
from docx.shared import Pt

f = open("demofile.docx", "xt")
document = Document("demofile.docx")
run = document.add_paragraph("piss").add_run()
font = run.font
font.name = 'Calibri'
font.size = Pt(65)
document.save('demo.docx')